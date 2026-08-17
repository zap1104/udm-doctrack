"""Getting text out of an uploaded file.

Two stages, in this order:

1. **Text layer** (free, offline, instant) — digital PDFs, Word files, plain
   text and spreadsheets already contain their text. Most office documents in a
   university are digital, so this handles the majority of uploads.
2. **OCR** (only when stage 1 finds nothing) — scanned pages and photos.
   Backends: OCR.space (free tier) and Azure Document Intelligence.

Everything degrades safely: if a library or an API key is missing, the document
is still saved and marked so it can be re-processed later.
"""

from __future__ import annotations

import io
import json
import logging
import time
import urllib.error
import urllib.request
from dataclasses import dataclass, field

from django.conf import settings

from apps.core.utils import file_extension, normalise_text

logger = logging.getLogger("doctrack")

TEXT_LIKE_EXTENSIONS = {"txt", "csv", "md", "log"}
IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "tif", "tiff", "bmp", "webp"}
MIN_USEFUL_CHARS = 40  # below this we treat a PDF as "scanned"


@dataclass
class ExtractionResult:
    text: str = ""
    engine: str = "none"
    pages: int = 0
    confidence: float | None = None
    status: str = "EMPTY"  # DONE | EMPTY | FAILED | SKIPPED
    notes: list[str] = field(default_factory=list)

    @property
    def char_count(self) -> int:
        return len(self.text)


# ---------------------------------------------------------------------------
# Stage 1 — text layer
# ---------------------------------------------------------------------------
def _extract_pdf(file_obj) -> ExtractionResult:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ExtractionResult(status="SKIPPED", notes=["pypdf is not installed"])

    try:
        file_obj.seek(0)
        reader = PdfReader(file_obj)
        pages = len(reader.pages)
        chunks = []
        for page in reader.pages:
            try:
                chunks.append(page.extract_text() or "")
            except Exception:  # a single broken page must not kill the upload
                continue
        text = normalise_text("\n".join(chunks))
        status = "DONE" if len(text) >= MIN_USEFUL_CHARS else "EMPTY"
        return ExtractionResult(text=text, engine="pdf-text-layer", pages=pages, status=status)
    except Exception as exc:
        logger.warning("PDF text extraction failed: %s", exc)
        return ExtractionResult(status="FAILED", notes=[str(exc)[:200]])


def _extract_docx(file_obj) -> ExtractionResult:
    try:
        import docx  # python-docx
    except ImportError:
        return ExtractionResult(status="SKIPPED", notes=["python-docx is not installed"])

    try:
        file_obj.seek(0)
        document = docx.Document(file_obj)
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        text = normalise_text("\n".join(parts))
        return ExtractionResult(
            text=text, engine="docx", pages=1, status="DONE" if text else "EMPTY"
        )
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        return ExtractionResult(status="FAILED", notes=[str(exc)[:200]])


def _extract_xlsx(file_obj) -> ExtractionResult:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ExtractionResult(status="SKIPPED", notes=["openpyxl is not installed"])

    try:
        file_obj.seek(0)
        workbook = load_workbook(file_obj, read_only=True, data_only=True)
        parts = []
        for sheet in workbook.worksheets:
            parts.append(sheet.title)
            for row in sheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value not in (None, "")]
                if values:
                    parts.append(" | ".join(values))
        workbook.close()
        text = normalise_text("\n".join(parts))
        return ExtractionResult(text=text, engine="xlsx", pages=1, status="DONE" if text else "EMPTY")
    except Exception as exc:
        logger.warning("XLSX extraction failed: %s", exc)
        return ExtractionResult(status="FAILED", notes=[str(exc)[:200]])


def _extract_plain(file_obj) -> ExtractionResult:
    try:
        file_obj.seek(0)
        raw = file_obj.read()
        text = normalise_text(raw.decode("utf-8", errors="replace") if isinstance(raw, bytes) else str(raw))
        return ExtractionResult(text=text, engine="plain-text", pages=1, status="DONE" if text else "EMPTY")
    except Exception as exc:
        return ExtractionResult(status="FAILED", notes=[str(exc)[:200]])


def extract_text_layer(file_obj, filename: str) -> ExtractionResult:
    extension = file_extension(filename)
    if extension == "pdf":
        return _extract_pdf(file_obj)
    if extension in {"docx", "dotx"}:
        return _extract_docx(file_obj)
    if extension in {"xlsx", "xlsm"}:
        return _extract_xlsx(file_obj)
    if extension in TEXT_LIKE_EXTENSIONS:
        return _extract_plain(file_obj)
    return ExtractionResult(status="EMPTY", engine="none", notes=[f"No text layer reader for .{extension}"])


# ---------------------------------------------------------------------------
# Stage 2 — OCR backends
# ---------------------------------------------------------------------------
def _prepare_ocr_payload(file_obj, filename: str) -> tuple[bytes, str, list[str]]:
    """Normalize image scans in memory; leave PDFs and unsupported images untouched."""
    file_obj.seek(0)
    payload = file_obj.read()
    extension = file_extension(filename)
    if extension not in IMAGE_EXTENSIONS:
        return payload, filename, []
    try:
        from PIL import Image, ImageEnhance, ImageOps

        with Image.open(io.BytesIO(payload)) as opened:
            image = ImageOps.exif_transpose(opened).convert("L")
            longest = max(image.size)
            if longest < 1600:
                scale = 1600 / max(longest, 1)
                image = image.resize(
                    (max(1, int(image.width * scale)), max(1, int(image.height * scale))),
                    Image.Resampling.LANCZOS,
                )
            image = ImageEnhance.Contrast(image).enhance(1.35)
            output = io.BytesIO()
            image.save(output, format="PNG", optimize=True)
            return output.getvalue(), f"{filename.rsplit('.', 1)[0]}.png", ["Image normalized for OCR."]
    except Exception as exc:
        logger.info("OCR image preprocessing skipped for %s: %s", filename, exc)
        return payload, filename, ["Image preprocessing was unavailable; original bytes were used."]


def _ocr_confidence(results: list[dict]) -> float | None:
    """Average optional word confidence values when the provider supplies them."""
    values: list[float] = []
    for result in results:
        overlay = result.get("TextOverlay") or {}
        for line in overlay.get("Lines") or []:
            for word in line.get("Words") or []:
                raw = word.get("Confidence")
                try:
                    value = float(raw)
                except (TypeError, ValueError):
                    continue
                values.append(value / 100 if value > 1 else value)
    return round(sum(values) / len(values), 4) if values else None


def _retryable_ocr_error(exc: Exception) -> bool:
    if isinstance(exc, urllib.error.HTTPError):
        return exc.code in {408, 425, 429} or 500 <= exc.code < 600
    return isinstance(exc, (urllib.error.URLError, TimeoutError))


def _ocr_space(file_obj, filename: str, *, language_hint: str = "auto") -> ExtractionResult:
    api_key = settings.OCR_SPACE_API_KEY
    if not api_key:
        return ExtractionResult(status="SKIPPED", notes=["OCR_SPACE_API_KEY is not set"])

    try:
        payload, provider_filename, notes = _prepare_ocr_payload(file_obj, filename)
        boundary = "----doctrackocrboundary"
        parts = []

        def add_field(name: str, value: str) -> None:
            parts.append(
                f"--{boundary}\r\nContent-Disposition: form-data; name=\"{name}\"\r\n\r\n{value}\r\n".encode()
            )

        provider_language = "eng"
        if language_hint == "fil":
            notes.append(
                "OCR.space has no Filipino locale mapping in this integration; its English model was used for Filipino/Taglish text."
            )
        add_field("apikey", api_key)
        add_field("language", provider_language)
        add_field("isOverlayRequired", "true")
        add_field("scale", "true")
        add_field("OCREngine", "2")
        parts.append(
            (
                f"--{boundary}\r\nContent-Disposition: form-data; name=\"file\"; filename=\"{provider_filename}\"\r\n"
                "Content-Type: application/octet-stream\r\n\r\n"
            ).encode()
        )
        parts.append(payload)
        parts.append(f"\r\n--{boundary}--\r\n".encode())
        body = b"".join(parts)
        retries = max(0, settings.OCR_PROVIDER_RETRIES)
        data = None
        for attempt in range(retries + 1):
            request = urllib.request.Request(
                settings.OCR_SPACE_ENDPOINT,
                data=body,
                headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
            )
            try:
                with urllib.request.urlopen(request, timeout=settings.OCR_PROVIDER_TIMEOUT_SECONDS) as response:
                    data = json.loads(response.read().decode("utf-8", errors="replace"))
                break
            except Exception as exc:
                if attempt >= retries or not _retryable_ocr_error(exc):
                    raise
                delay = settings.OCR_RETRY_BASE_SECONDS * (2**attempt)
                notes.append(f"OCR provider retry {attempt + 1} after a temporary error.")
                logger.warning("OCR.space transient failure; retrying in %ss: %s", delay, exc)
                time.sleep(delay)

        if data is None:
            raise ValueError("OCR.space returned no response")
        if data.get("IsErroredOnProcessing"):
            message = "; ".join(data.get("ErrorMessage") or ["OCR.space reported an error"])
            return ExtractionResult(status="FAILED", engine="ocr.space", notes=notes + [message[:200]])

        results = data.get("ParsedResults") or []
        text = normalise_text("\n".join(item.get("ParsedText", "") for item in results))
        confidence = _ocr_confidence(results)
        if confidence is None:
            notes.append("The OCR provider did not return confidence values.")
        return ExtractionResult(
            text=text,
            engine=f"ocr.space:{provider_language}",
            pages=len(results),
            confidence=confidence,
            status="DONE" if text else "EMPTY",
            notes=notes,
        )
    except (urllib.error.URLError, TimeoutError, ValueError) as exc:
        logger.warning("OCR.space call failed: %s", exc)
        return ExtractionResult(status="FAILED", engine="ocr.space", notes=[str(exc)[:200]])


def _ocr_azure(file_obj, filename: str, *, language_hint: str = "auto") -> ExtractionResult:
    endpoint = settings.AZURE_DOCINT_ENDPOINT
    key = settings.AZURE_DOCINT_KEY
    if not (endpoint and key):
        return ExtractionResult(status="SKIPPED", notes=["Azure Document Intelligence is not configured"])
    try:
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.core.credentials import AzureKeyCredential
        from azure.core.exceptions import HttpResponseError, ServiceRequestError, ServiceResponseError
    except ImportError:
        return ExtractionResult(status="SKIPPED", notes=["azure-ai-documentintelligence is not installed"])

    try:
        payload, _provider_filename, notes = _prepare_ocr_payload(file_obj, filename)
        client = DocumentIntelligenceClient(endpoint=endpoint, credential=AzureKeyCredential(key))
        retries = max(0, int(settings.OCR_PROVIDER_RETRIES))
        locale = {"eng": "en-US", "fil": "fil-PH"}.get(language_hint)
        result = None
        for attempt in range(retries + 1):
            try:
                options = {"locale": locale} if locale else {}
                poller = client.begin_analyze_document("prebuilt-read", body=payload, **options)
                result = poller.result()
                break
            except (ServiceRequestError, ServiceResponseError, HttpResponseError) as exc:
                status_code = getattr(exc, "status_code", None)
                transient = not isinstance(exc, HttpResponseError) or status_code == 429 or (
                    status_code is not None and status_code >= 500
                )
                if not transient or attempt >= retries:
                    raise
                delay = max(0, int(settings.OCR_RETRY_BASE_SECONDS)) * (2**attempt)
                notes.append(f"Azure OCR transient failure; retry {attempt + 1} of {retries} after {delay}s.")
                logger.warning("Azure OCR transient failure; retrying in %ss: %s", delay, exc)
                time.sleep(delay)
        if result is None:
            raise ValueError("Azure Document Intelligence returned no response")
        text = normalise_text(getattr(result, "content", "") or "")
        pages = getattr(result, "pages", []) or []
        confidence_values = [
            float(word.confidence)
            for page in pages
            for word in (getattr(page, "words", None) or [])
            if getattr(word, "confidence", None) is not None
        ]
        confidence = round(sum(confidence_values) / len(confidence_values), 4) if confidence_values else None
        if confidence is None:
            notes.append("The OCR provider did not return confidence values.")
        return ExtractionResult(
            text=text,
            engine=f"azure-document-intelligence:{language_hint}",
            pages=len(pages),
            confidence=confidence,
            status="DONE" if text else "EMPTY",
            notes=notes,
        )
    except Exception as exc:  # pragma: no cover - depends on external service
        logger.warning("Azure Document Intelligence failed: %s", exc)
        return ExtractionResult(status="FAILED", engine="azure-document-intelligence", notes=[str(exc)[:200]])


def run_ocr(
    file_obj, filename: str, *, language_hint: str = "auto", allow_external_ocr: bool = True
) -> ExtractionResult:
    backend = (settings.OCR_BACKEND or "auto").lower()
    if not allow_external_ocr:
        return ExtractionResult(
            status="SKIPPED",
            notes=["External OCR was disabled for this sensitive document; no file content was sent to a provider."],
        )
    if backend == "none":
        return ExtractionResult(status="SKIPPED", notes=["OCR disabled by configuration"])
    if backend == "azure":
        return _ocr_azure(file_obj, filename, language_hint=language_hint)
    if backend == "ocrspace":
        return _ocr_space(file_obj, filename, language_hint=language_hint)
    # auto: try whichever is configured, cheapest first
    if settings.OCR_SPACE_API_KEY:
        result = _ocr_space(file_obj, filename, language_hint=language_hint)
        if result.status == "DONE":
            return result
    if settings.AZURE_DOCINT_ENDPOINT and settings.AZURE_DOCINT_KEY:
        return _ocr_azure(file_obj, filename, language_hint=language_hint)
    return ExtractionResult(status="SKIPPED", notes=["No OCR provider configured"])


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------
def extract_document_text(
    file_obj, filename: str, *, language_hint: str = "auto", allow_external_ocr: bool = True
) -> ExtractionResult:
    """Text layer first, OCR only if needed. Never raises."""
    extension = file_extension(filename)
    try:
        if extension in IMAGE_EXTENSIONS:
            result = run_ocr(
                file_obj, filename, language_hint=language_hint, allow_external_ocr=allow_external_ocr
            )
            if result.status != "DONE" and allow_external_ocr:
                result.notes.append("Image upload needs OCR; configure OCR_SPACE_API_KEY to enable it.")
            return _clip(result)

        layer = extract_text_layer(file_obj, filename)
        if layer.status == "DONE" and len(layer.text) >= MIN_USEFUL_CHARS:
            return _clip(layer)

        ocr = run_ocr(
            file_obj, filename, language_hint=language_hint, allow_external_ocr=allow_external_ocr
        )
        if ocr.status == "DONE":
            ocr.notes.append("Text layer was empty, so OCR was used.")
            ocr.pages = ocr.pages or layer.pages
            return _clip(ocr)

        layer.notes.extend(ocr.notes)
        if layer.status == "EMPTY":
            layer.notes.append("No text found. Metadata can still be entered by hand.")
        return _clip(layer)
    except Exception as exc:  # pragma: no cover - defensive
        logger.exception("Text extraction crashed for %s", filename)
        return ExtractionResult(status="FAILED", notes=[str(exc)[:200]])


def _clip(result: ExtractionResult) -> ExtractionResult:
    limit = settings.OCR_MAX_CHARS
    if len(result.text) > limit:
        result.text = result.text[:limit]
        result.notes.append(f"Text truncated to {limit} characters.")
    return result
